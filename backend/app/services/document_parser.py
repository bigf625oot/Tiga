import io
from fastapi import UploadFile, HTTPException
import pypdf
import docx
import os
import logging
from app.services.markdown import to_markdown
from app.core.config import settings
from typing import Optional
from PIL import Image
import base64

logger = logging.getLogger(__name__)

def parse_local_file(file_path: str) -> str:
    """
    Parse a local file and return text content.
    """
    filename = os.path.basename(file_path).lower()
    logger.info(f"Parsing local file: {file_path}")
    text = ""
    
    try:
        if filename.endswith('.pdf'):
            logger.info("Parsing PDF file...")
            parsed = False
            # 1) PyMuPDF
            try:
                import fitz
                doc = fitz.open(file_path)
                try:
                    parts = []
                    for i in range(doc.page_count):
                        p = doc.load_page(i)
                        parts.append(p.get_text("text") or "")
                    text = "\n".join(parts)
                    parsed = True
                    logger.info(f"PyMuPDF parsed successfully. Pages: {doc.page_count}")
                finally:
                    doc.close()
            except Exception as e1:
                logger.warning(f"PyMuPDF parse failed: {e1}")
            # 2) pdfplumber
            if not parsed:
                try:
                    import pdfplumber
                    parts = []
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            parts.append(page.extract_text() or "")
                    text = "\n".join(parts)
                    parsed = True
                    logger.info(f"pdfplumber parsed successfully. Pages: {len(parts)}")
                except Exception as e2:
                    logger.warning(f"pdfplumber parse failed: {e2}")
            # 3) pypdf
            if not parsed:
                try:
                    with open(file_path, "rb") as f:
                        reader = pypdf.PdfReader(f)
                        parts = []
                        for i, page in enumerate(reader.pages):
                            try:
                                parts.append(page.extract_text() or "")
                            except Exception:
                                parts.append("")
                        text = "\n".join(parts)
                    parsed = True
                    logger.info(f"pypdf parsed successfully. Pages: {len(parts)}")
                except Exception as e3:
                    logger.warning(f"pypdf parse failed: {e3}")
            if not parsed:
                raise RuntimeError("PDF parsing failed on all strategies")
            if settings and getattr(settings, "OCR_ENABLED", False) and not (text or "").strip():
                try:
                    import fitz, pytesseract
                    doc = fitz.open(file_path)
                    try:
                        parts = []
                        for i in range(doc.page_count):
                            p = doc.load_page(i)
                            pm = p.get_pixmap()
                            img_bytes = pm.tobytes("png")
                            img = Image.open(io.BytesIO(img_bytes))
                            parts.append(pytesseract.image_to_string(img))
                        text = "\n".join(parts)
                        logger.info(f"OCR parsed successfully. Pages: {doc.page_count}")
                    finally:
                        doc.close()
                except Exception as eocr:
                    logger.warning(f"OCR parse failed: {eocr}")
        elif filename.endswith('.docx'):
            logger.info("Parsing DOCX file...")
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            logger.info(f"DOCX parsed successfully. Paragraphs: {len(doc.paragraphs)}")
        else:
            # Try plain text
            logger.info("Parsing plain text file...")
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            except:
                try:
                    with open(file_path, "r", encoding="gbk") as f:
                        text = f.read()
                except:
                    pass
            logger.info("Plain text parsed.")
                    
        text = sanitize_text(text.strip())
        text = to_markdown(text, {"source": "local", "title": os.path.splitext(os.path.basename(file_path))[0]})
        return text
    except Exception as e:
        logger.error(f"Error parsing local file {file_path}: {e}")
        print(f"Error parsing local file {file_path}: {e}")
        return ""

def sanitize_text(text: str) -> str:
    """
    清洗文本，移除可能导致编码问题的非法字符（如单独的代理对）。
    """
    if not text:
        return ""
    # 1. 移除非法代理对 (surrogates)
    # 使用 errors='ignore' 忽略无法编码的字符，然后再解码回来
    cleaned = text.encode('utf-8', 'ignore').decode('utf-8')
    
    # 2. 移除空字符 (NUL) - 某些数据库如 Postgres 不支持
    cleaned = cleaned.replace('\x00', '')
    
    return cleaned

async def parse_document(file: UploadFile) -> str:
    """
    Parse uploaded document (PDF or Word) and return text content.
    """
    content_type = file.content_type
    filename = file.filename.lower()
    
    try:
        content = await file.read()
        file_stream = io.BytesIO(content)
        
        text = ""
        
        if filename.endswith('.pdf') or content_type == 'application/pdf':
            parsed = False
            # 1) PyMuPDF from bytes
            try:
                import fitz
                doc = fitz.open(stream=content, filetype="pdf")
                try:
                    parts = []
                    for i in range(doc.page_count):
                        p = doc.load_page(i)
                        parts.append(p.get_text("text") or "")
                    text = "\n".join(parts)
                    parsed = True
                finally:
                    doc.close()
            except Exception as e1:
                logger.warning(f"PyMuPDF bytes parse failed: {e1}")
            # 2) pdfplumber from BytesIO
            if not parsed:
                try:
                    import pdfplumber
                    parts = []
                    with pdfplumber.open(file_stream) as pdf:
                        for page in pdf.pages:
                            parts.append(page.extract_text() or "")
                    text = "\n".join(parts)
                    parsed = True
                except Exception as e2:
                    logger.warning(f"pdfplumber bytes parse failed: {e2}")
            # 3) pypdf
            if not parsed:
                try:
                    reader = pypdf.PdfReader(file_stream)
                    for page in reader.pages:
                        try:
                            text += (page.extract_text() or "") + "\n"
                        except Exception:
                            text += "\n"
                    parsed = True
                except Exception as e3:
                    logger.warning(f"pypdf bytes parse failed: {e3}")
            if not text.strip():
                if settings and getattr(settings, "OCR_ENABLED", False):
                    try:
                        import fitz, pytesseract
                        doc = fitz.open(stream=content, filetype="pdf")
                        try:
                            parts = []
                            for i in range(doc.page_count):
                                p = doc.load_page(i)
                                pm = p.get_pixmap()
                                img_bytes = pm.tobytes("png")
                                img = Image.open(io.BytesIO(img_bytes))
                                parts.append(pytesseract.image_to_string(img))
                            text = "\n".join(parts)
                        finally:
                            doc.close()
                    except Exception as eocr:
                        logger.warning(f"OCR bytes parse failed: {eocr}")
                if not text.strip():
                    raise HTTPException(status_code=400, detail="未提取到文本内容。如果是扫描版 PDF，请先使用 OCR 工具或开启 OCR_ENABLED。")
                
        elif filename.endswith(('.docx', '.doc')) or 'word' in content_type:
            # python-docx only supports .docx
            if filename.endswith('.doc'):
                raise HTTPException(status_code=400, detail="暂不支持旧版 .doc 格式，请另存为 .docx 后上传")
                
            doc = docx.Document(file_stream)
            for para in doc.paragraphs:
                text += para.text + "\n"
                
            if not text.strip():
                raise HTTPException(status_code=400, detail="未提取到文本内容。文档可能只包含图片或表格。")
        else:
            # Try to read as plain text
            try:
                text = content.decode('utf-8')
            except:
                # Try gbk
                try:
                    text = content.decode('gbk')
                except:
                    raise HTTPException(status_code=400, detail="无法解码文件内容，请确保是 UTF-8 或 GBK 编码的文本文件。")
                
        text = sanitize_text(text.strip())
        text = to_markdown(text, {"source": "upload", "title": os.path.splitext(file.filename)[0]})
        return text
        
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        print(f"Error parsing document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")
